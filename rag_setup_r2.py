import os
import glob
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores.faiss import FAISS
from tqdm import tqdm
import pandas as pd
from langchain.schema import Document

# txt/md 파일 인코딩 자동 감지
def load_text_any_encoding(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text)]
    except Exception:
        try:
            with open(path, "r", encoding="euc-kr") as f:
                text = f.read()
            return [Document(page_content=text)]
        except Exception as e:
            print(f"[Error] TXT/MD 파일 파싱 실패: {path} ({e})")
            return []

# csv 파일 파싱
def load_csv(path):
    try:
        df = pd.read_csv(path, encoding="utf-8")
        text = df.to_string(index=False)
        return [Document(page_content=text)]
    except Exception:
        try:
            df = pd.read_csv(path, encoding="euc-kr")
            text = df.to_string(index=False)
            return [Document(page_content=text)]
        except Exception as e:
            print(f"[Error] CSV 파일 파싱 실패: {path} ({e})")
            return []

# xlsx 파일 파싱
def load_xlsx(path):
    try:
        df = pd.read_excel(path)
        text = df.to_string(index=False)
        return [Document(page_content=text)]
    except Exception as e:
        print(f"[Error] XLSX 파일 파싱 실패: {path} ({e})")
        return []

# docx/doc 파일 파싱
def load_docx(path):
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return [Document(page_content=text)]
    except Exception as e:
        print(f"[Error] DOCX/DOC 파일 파싱 실패: {path} ({e})")
        return []

# 환경변수 로드
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

# 1. 파일 전체 재귀 탐색 및 로딩
target_folders = [
    "docs_r2/금융용어 설명 자료",
    "docs_r2/은행내규",
    "docs_r2/QnA",
    "docs_r2/상품정보"
]
docs = []
for folder in target_folders:
    for fpath in glob.glob(folder + '/**/*', recursive=True):
        if not os.path.isfile(fpath):
            continue
        if os.path.getsize(fpath) == 0:
            print(f"[Warning] 파일이 비어 있어 스킵: {fpath}")
            continue
        fname = os.path.basename(fpath)
        ext = fname.lower().split('.')[-1]
        try:
            if ext == "pdf":
                try:
                    docs.extend(PyPDFLoader(fpath).load())
                except Exception as e:
                    print(f"[Error] PDF 파일 파싱 실패: {fpath} ({e})")
            elif ext in ["txt", "md"]:
                docs.extend(load_text_any_encoding(fpath))
            elif ext in ["docx", "doc"]:
                docs.extend(load_docx(fpath))
            elif ext == "csv":
                docs.extend(load_csv(fpath))
            elif ext == "xlsx":
                docs.extend(load_xlsx(fpath))
            # hwp 등 기타 포맷은 아예 무시 (pass)
        except Exception as e:
            print(f"[Error] 파일 파싱 중 알 수 없는 에러: {fpath} ({e})")
            continue

print(f"실제 파일 로드 수: {len(docs)}")
for i, d in enumerate(docs[:3]):
    print(f"문서 {i+1} 미리보기: {getattr(d, 'page_content', '')[:100]}")

# 2. 텍스트 청킹
splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200
)
chunks = splitter.split_documents(docs)

# 3. 임베딩 생성 (배치 처리)
embeddings = OpenAIEmbeddings(openai_api_key=API_KEY)
all_embeddings = []
batch_size = 100

print(f"총 {len(chunks)}개 청크를 {batch_size}개씩 임베딩합니다...")

for i in tqdm(range(0, len(chunks), batch_size)):
    batch = chunks[i:i+batch_size]
    texts = [doc.page_content for doc in batch]
    vectors = embeddings.embed_documents(texts)
    all_embeddings.extend(vectors)

# 4. FAISS 벡터 DB 구축 및 저장
text_embedding_pairs = list(zip([doc.page_content for doc in chunks], all_embeddings))
if not text_embedding_pairs:
    print("❌ 벡터 DB에 저장할 데이터가 없습니다. 문서가 제대로 로드되었는지 확인하세요.")
else:
    vectorstore = FAISS.from_embeddings(
        text_embedding_pairs,
        embeddings
    )
    vectorstore.save_local("faiss_index_full")
    print("✅ RAG 벡터 인덱스 생성 완료: faiss_index_full 폴더에 저장됨")
